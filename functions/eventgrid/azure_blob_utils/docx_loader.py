import tempfile
from docx import Document as DocxDocument
import logging
from typing import List
from langchain_core.documents import Document
from .base_loader import BaseDocumentLoader, DocumentLoaderFactory

class DocxLoader(BaseDocumentLoader):
    def load(self, blob_name: str) -> List[Document]:
        container_client = self.blob_service_client.get_container_client(self.container_name)
        blob_client = container_client.get_blob_client(blob_name)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            blob_data = blob_client.download_blob()
            blob_data.readinto(temp_file)
            temp_path = temp_file.name
            
        try:
            doc = DocxDocument(temp_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            if text.strip():
                metadata = {
                    'source': f"{self.container_name}/{blob_name}",
                    'file_type': 'docx'
                }
                return [Document(page_content=text, metadata=metadata)]
            return []
            
        except Exception as e:
            logging.error(f"Error processing DOCX {blob_name}: {str(e)}")
            raise
        finally:
            import os
            os.unlink(temp_path)

# Register the loader
DocumentLoaderFactory.register_loader('.docx', DocxLoader)
DocumentLoaderFactory.register_loader('.doc', DocxLoader)